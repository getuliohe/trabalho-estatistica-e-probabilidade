from math import sqrt
import numpy as np
from sklearn.linear_model import LinearRegression
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

class Analise:

    def __init__(self, lista, listax, listay) -> None:
        self.n = len(lista)
        self.media = Analise.calculaMedia(lista)
        self.minimo = min(lista)
        self.quartil1 = Analise.calcularQuadrantes(lista)[0]
        self.quartil2 = Analise.calcularQuadrantes(lista)[1]
        self.quartil3 = Analise.calcularQuadrantes(lista)[2]
        self.maximo = max(lista)
        self.mediana = Analise.calcularMediana(lista)
        self.varancia = Analise.calcularVariancia(lista)
        self.desvioPadrao = Analise.calcularDesvioPadrao(lista)
        self.desvioMedioAbsoluto = Analise.calcularDesvioMedioAbsoluto(lista)
        self.amplitude = Analise.calcularAmplitude(lista)
        self.total = sum(lista)
        self.coeficienteAngular = self.regressaoLinear(listax, listay)[0]
        self.coeficienteLinear = self.regressaoLinear(listax, listay)[1]

    @staticmethod
    def calculaMedia(lista):
        media = 0
        listaFiltrada = lista
        totalLinhas = len(listaFiltrada)

        for linha in listaFiltrada:
                media += linha

        media /= totalLinhas
        return media
    
    @staticmethod
    def calcularQuadrantes(lista):
        listaEmRol = sorted(filter(None, lista))
        
        def calcularQ1(lista):
            n = len(lista)
            if n % 2 == 0:
                return (lista[n//4 - 1] + lista[n//4]) / 2
            else:
                return lista[n//4]

        def calcularQ2(lista):
            n = len(lista)
            if n % 2 == 0:
                return (lista[n//2 - 1] + lista[n//2]) / 2
            else:
                return lista[n//2]

        def calcularQ3(lista):
            n = len(lista)
            if n % 2 == 0:
                return (lista[(3*n)//4 - 1] + lista[(3*n)//4]) / 2
            else:
                return lista[(3*n)//4]
            
        return calcularQ1(listaEmRol), calcularQ2(listaEmRol), calcularQ3(listaEmRol)

    
    @staticmethod
    def calcularMediana(lista):
        listaEmRol = sorted(lista)
        n = len(listaEmRol)

        if n % 2 != 0:
            return listaEmRol[(n // 2)]
        else:
            return (listaEmRol[(n // 2) - 1] + listaEmRol[(n // 2)])/2
    
    @staticmethod
    def calcularVariancia(lista):
        listaFiltrada = lista
        media = Analise.calculaMedia(lista)
        total = 0

        for item in listaFiltrada:
            total += (item - media)**2
        
        return total/(len(listaFiltrada)- 1)
    
    @staticmethod
    def calcularDesvioPadrao(lista):
        return sqrt(Analise.calcularVariancia(lista))
    
    @staticmethod
    def calcularDesvioMedioAbsoluto(lista):
        listaFiltrada = lista
        media = Analise.calculaMedia(lista)
        total = [abs(x - media) for x in listaFiltrada]
        
        return sum(total)/len(listaFiltrada)
    
    @staticmethod
    def calcularAmplitude(lista):
        return (max(lista) - min(lista))
    
    @staticmethod
    def calcularCoeficienteDeVaricao(lista):
        media = Analise.calculaMedia(lista)

        return (Analise.calcularDesvioPadrao(lista)/ media) * 100
    
    @staticmethod
    def filtrarNulos(lista):
        listaSemNulos = []

        for item in lista:
            if item != None or item == 0:
                listaSemNulos.append(item)

        return listaSemNulos
    
    def regressaoLinear(self,listax,listay):
        X = np.array(listax)
        y = np.array(listay)

        model = LinearRegression()

        model.fit(X, y)

        coeficiente_angular = model.coef_[0]
        coeficiente_linear = model.intercept_

        return coeficiente_angular, coeficiente_linear
    
    def graficoRegressaoLinear(self, listax, listay):
        X = np.array(listax)
        y = np.array(listay)


        model = LinearRegression()

        model.fit(X, y)

        y_pred = model.predict(X)
        
        plt.scatter(X, y, color='blue', label='Dados reais')
        plt.plot(X, y_pred, color='red', label='Linha de regressão')
        plt.xlabel('X')
        plt.ylabel('y')
        plt.legend()

        plot_path = f"Regressao linear População censo 2022.png"
        plt.savefig(plot_path)
        plt.close()

        return plot_path

    def gerarGraficoBoxPLot(lista):

        plt.figure(figsize=(8,6))
        plt.boxplot(lista, patch_artist=True, showmeans=True, showfliers=True)
        plt.title("Bloxplox da media")
        plt.ylabel('Valores')
        plt.grid(True)
        plt.yticks(range(min(lista),max(lista),30000000))

        plot_path = f"boxPlotPopulação censo 2022.png"
        plt.savefig(plot_path)
        plt.close()

        return plot_path
    
    def gerarDocumento(self,pathsGraficos):
        doc = Document()

        doc.add_heading(f"Analise do censo populacional 2022")

        texto = f"n = {self.n}\nmédia = {self.media}\nmediana = {self.mediana}\nmínimo = {self.minimo}\nquartil 1 = {self.quartil1}\nquartil 2 = {self.quartil2}\nquartil 3 = {self.quartil3}\nmáximo = {self.maximo}\nvariancia = {self.varancia}\ndesvio padrão = {self.desvioPadrao}\n desvio médio absoluto = {self.desvioMedioAbsoluto}\namplitude = {self.amplitude}\ntotal = {self.total}\ncoficiente angular = {self.coeficienteAngular}\ncoeficiente linear = {self.coeficienteLinear}"

        doc.add_paragraph(texto)

        for path in pathsGraficos:
            doc.add_heading(path, level=1)
            doc.add_picture(path, width=Inches(6))
        
        doc.save(f"Analise do censo populacional 2022.docx")
 